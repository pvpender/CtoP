from docx import Document
import re


class Parser:
    def __init__(self, text: str):
        text = text.replace("int main()", 'Алг "Лабораторная работа"')
        self.strings = list(filter(lambda x: ("#include" not in x) and (x != ""), text.split("\n")))
        self.parse_cycles()
        self.parse_input_out()
        self.replace_while()
        self.parse_if()
        self.strings = [self.strings[i].replace("}", "кц").removesuffix(";").replace("++", "+= 1").replace("--", "-= 1")
                        for i in range(len(self.strings))]
        print(self.strings)

    def parse(self):
        for i in range(len(self.strings)):
            if re.search(r'\w', self.strings[i]):
                print(self.strings[i])

    def parse_cycles(self):
        for i in range(len(self.strings)):
            if "for" in self.strings[i]:
                mas = [re.findall(r'[\d \w+-]+', j)[0].replace(" ", "")
                       for j in re.findall(r'[=<>][\w \d+-]+', self.strings[i])]
                s = re.findall(r'[\d \w]+=', self.strings[i])[0]
                var_name = s.replace(" ", "").replace("=", "") if s.count(" ") <= 1 \
                    else s[s.find(" "):s.find("=")].replace(" ", "")
                self.strings[i] = self.strings[i][0:self.strings[i].find("f")] + f"цикл от {var_name}:={mas[0]} до " \
                                                                                 f"{mas[1]}"
            if "while" in self.strings[i]:
                self.strings[i] = self.strings[i].replace("while", "цикл пока").replace(";", "").removesuffix(";").\
                    removesuffix("{")

    def parse_input_out(self):
        for i in range(len(self.strings)):
            if "printf" in self.strings[i]:
                if re.search(r'^(?!.*\",).*$', self.strings[i].replace(" ", "")):
                    text = re.findall(r'(?<=\()([\s\S]+?)(?=\))', self.strings[i].replace("\\n", ""))[0]
                    self.strings[i] = self.strings[i][0:self.strings[i].find("p")] + f"вывод({text})"
                else:
                    a = self.strings[i][self.strings[i].rfind('"') + 2:].removesuffix(";").removesuffix(")")
                    mas_variables = re.findall(r'[A-Za-z0-9_]+\([A-Za-z0-9_,\s\[\]]+\)|[A-Za-z\[\]_0-9-\s]+', a)
                    # print(re.findall(r'[\"\'][\S\s]+[\"\']', self.strings[i]))
                    mas_pattern = re.findall(r'%[\w]+', self.strings[i])
                    for j in range(len(mas_variables)):
                        print(mas_pattern, mas_variables)
                        self.strings[i] = self.strings[i].replace(mas_pattern[j], mas_variables[j], 1)
                    # print(self.strings[i].replace("%", "").replace("\\n", ""))
                    self.strings[i] = (self.strings[i].replace("printf_s", "вывод") if "printf_s" in self.strings[i]
                                       else self.strings[i].replace("printf", "вывод")).removesuffix(";").replace("\\n",
                                                                                                                  "")
            elif "scanf" in self.strings[i]:
                self.strings[i] = self.strings[i][:self.strings[i].find('"')] + \
                                  self.strings[i][self.strings[i].rfind('"') + 3:]
                self.strings[i] = (self.strings[i].replace("scanf_s", "ввод") if "printf_s" in self.strings[i]
                                   else self.strings[i].replace("scanf", "ввод")).removesuffix(";").replace("\\n", "").\
                    replace("&", "")

    def replace_while(self):
        mas_do = []
        for i in range(len(self.strings)):
            if "do" in self.strings[i]:
                mas_do.append(i)
            elif "цикл пока" in self.strings[i] and "}" in self.strings[i] and mas_do:
                self.strings[mas_do[-1]] = self.strings[mas_do[-1]][:self.strings[mas_do[-1]].find("d")] + \
                                           self.strings[i][self.strings[i].find("ц"):]
                self.strings[mas_do[-1] + 1] = self.strings[mas_do[-1] + 1].replace("{", "")
                self.strings[i] = (self.strings[i][:self.strings[i].find("ц")] + "кц").replace("}", "")
                del mas_do[-1]

    def parse_if(self):
        mas_if = []
        mas_else = []
        for i in range(len(self.strings)):
            if "if" in self.strings[i]:
                mas_if.append(i)
                self.strings[i] = f'{self.strings[i][:self.strings[i].find("i")]}если ' \
                                  f'{self.strings[i][self.strings[i].find("(") + 1 : self.strings[i].find(")")-1]} тo'
            elif "цикл" in self.strings[i]:
                mas_if.append(i)
            elif "}" in self.strings[i] and mas_if:
                if "цикл" in self.strings[mas_if[-1]]:
                    del mas_if[-1]
                else:
                    del mas_if[-1]
                    if "else" in self.strings[i+1]:
                        self.strings[i] = self.strings[i].replace("}", "")
                    elif "else" in self.strings[i]:
                        self.strings[i] = self.strings[i].replace("}", "", 1)
                    else:
                        self.strings[i] = self.strings[i].replace("}", "всё")
            if "else" in self.strings[i]:
                mas_else.append(i)
                self.strings[i] = self.strings[i].replace("else", "иначе").replace("{", "")
            elif "цикл" in self.strings[i]:
                mas_else.append(i)
            elif "}" in self.strings[i] and mas_else:
                if "цикл" in self.strings[mas_else[-1]]:
                    del mas_else[-1]
                else:
                    del mas_else[-1]
                    self.strings[i] = self.strings[i].replace("}", "всё")


class Writer:
    def __init__(self, strings, file, list_keywords, path):
        doc = Document()
        for i in range(len(strings)):
            file.write(f"{strings[i]}\n")
            p = doc.add_paragraph(strings[i][:(len(strings[i]) - len(strings[i].lstrip()))])
            mas = strings[i][len(strings[i]) - len(strings[i].lstrip()):].split(" ")
            for j in mas:
                if any(ele == j for ele in list_keywords):
                    p.add_run(f"{j}").underline = True
                    p.add_run(" ").underline = False
                else:
                    p.add_run(f"{j} ").underline = False
        doc.save(path)
